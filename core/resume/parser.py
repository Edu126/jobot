"""Parse uploaded resume (DOCX or PDF) into a structured dict.

Output schema:
    {
        "source_format": "docx" | "pdf",
        "raw_text": str,
        "contact": {"name", "email", "phone", "location", "linkedin", "website"},
        "sections": {section_name: list[str]},   # paragraphs/bullets per section
        "stats": {"word_count": int, "page_estimate": int, "bullet_count": int},
        "docx_meta": {"has_tables": bool, "has_images": bool, "table_count": int, "image_count": int},
    }

Parsing is best-effort. Real resumes vary wildly in layout; this module
extracts what it can and leaves the rest as raw text.
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from pypdf import PdfReader


SECTION_PATTERNS: dict[str, list[str]] = {
    "summary": [
        r"^\s*(summary|profile|professional\s+summary|career\s+summary|career\s+profile|"
        r"career\s+objective|objective|about(\s+me)?|overview|executive\s+summary|"
        r"personal\s+statement|highlights)\s*:?\s*$",
    ],
    "experience": [
        r"^\s*(experience|work\s+experience|professional\s+experience|relevant\s+experience|"
        r"employment(\s+history)?|career(\s+history)?|career\s+experience|work\s+history|"
        r"professional\s+background|career\s+background|background|"
        r"professional\s+history|relevant\s+work\s+experience|project\s+experience)\s*:?\s*$",
    ],
    "education": [
        r"^\s*(education|education\s*&?\s*training|academic(\s+background|\s+qualifications|\s+credentials|\s+history)?|"
        r"qualifications|degrees|educational\s+background|academic|"
        r"studies|formation|formation\s+académique)\s*:?\s*$",
    ],
    "skills": [
        r"^\s*(skills|technical\s+skills|core\s+competencies|core\s+skills|key\s+skills|"
        r"competencies|areas\s+of\s+expertise|expertise|professional\s+skills|"
        r"tools(\s*&?\s*technologies)?|technologies|software(\s+proficiency)?|"
        r"technical\s+proficiency|technical\s+tools|software\s*&?\s*tools|"
        r"hard\s+skills|soft\s+skills|computer\s+skills|relevant\s+skills)\s*:?\s*$",
    ],
    "certifications": [
        r"^\s*(certifications?|licenses?(\s*&?\s*certifications?)?|certificates?|"
        r"professional\s+certifications?|accreditations?|"
        r"professional\s+development|continuing\s+education)\s*:?\s*$",
    ],
    "projects": [
        r"^\s*(projects?|key\s+projects?|notable\s+projects?|selected\s+projects?|"
        r"project\s+highlights|portfolio)\s*:?\s*$",
    ],
    "volunteer": [
        r"^\s*(volunteer(ing)?|volunteer\s+experience|volunteer\s+work|"
        r"community(\s+involvement|\s+service)?|civic\s+engagement)\s*:?\s*$",
    ],
    "awards": [
        r"^\s*(awards?|honors?|honou?rs?(\s*&?\s*awards?)?|achievements?|"
        r"accomplishments?|distinctions?|recognition)\s*:?\s*$",
    ],
    "languages": [
        r"^\s*(languages?|language\s+proficiency|spoken\s+languages?)\s*:?\s*$",
    ],
    "interests": [
        r"^\s*(interests?|hobbies(\s*&?\s*interests?)?|personal\s+interests?)\s*:?\s*$",
    ],
    "references": [
        r"^\s*(references?|professional\s+references?)\s*:?\s*$",
    ],
}

EMAIL_RE = re.compile(r"[\w._%+-]+@[\w.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(
    r"(\+?\d{1,2}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}"
)
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-%]+", re.I)
URL_RE = re.compile(r"https?://[\w./\-_%?=&#]+", re.I)
# Loose Ontario-city hint; expand later if needed.
LOCATION_RE = re.compile(
    r"\b(Ottawa|Gatineau|Kanata|Nepean|Orleans|Toronto|Montreal|Vancouver|Calgary|Edmonton|Winnipeg|Halifax|Quebec)\b[^\n|]*",
    re.I,
)


def parse_resume(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """Entry point. Dispatches on filename extension."""
    name = filename.lower()
    if name.endswith(".docx"):
        return _parse_docx(file_bytes)
    if name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    raise ValueError(f"Unsupported file type: {filename}. Use .docx or .pdf.")


def _parse_docx(data: bytes) -> dict[str, Any]:
    doc = Document(io.BytesIO(data))

    # Collect text per paragraph in document order. We do NOT prefix bullets
    # with '•' here — list membership is conveyed by the line existing inside
    # a section (the writer turns each section line back into a Word bullet),
    # and prefixing here causes double-bullets in display and confuses the LLM.
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Strip any bullet glyphs the source document hard-coded as text
        # rather than as a real list style.
        text = _strip_leading_bullet(text)
        if text:
            lines.append(text)

    # Include table cell text too — many resumes use tables for layout.
    table_count = len(doc.tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)

    image_count = _count_docx_images(doc)
    raw_text = "\n".join(lines)

    parsed = _common_parse(raw_text)
    parsed["source_format"] = "docx"
    parsed["docx_meta"] = {
        "has_tables": table_count > 0,
        "has_images": image_count > 0,
        "table_count": table_count,
        "image_count": image_count,
    }
    return parsed


def _count_docx_images(doc: Document) -> int:
    count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            count += 1
    return count


# Bullet characters that signal the START of a new line — never merge INTO
# a previous line if the next starts with one of these.
_BULLET_CHARS = ("•", "·", "●", "○", "◆", "◇", "■", "□", "▪", "▫", "★", "☆", "-", "*", "–", "—")

# Common resume section headers — line starting with any of these is a
# block boundary, don't merge into the previous line.
_SECTION_HEADER_HINTS = (
    "profile", "summary", "objective", "experience", "work experience",
    "employment", "education", "skills", "technical skills", "projects",
    "certifications", "certificates", "licenses", "volunteer",
    "volunteering", "awards", "publications", "languages", "interests",
    "hobbies", "references", "additional information", "achievements",
)


def _looks_like_new_block(line: str) -> bool:
    """True when a line clearly starts its own block — bullet, section
    header, or a role/company line with a year in it. Used by the reflow
    heuristic to KNOW not to merge into whatever came before."""
    s = line.strip()
    if not s:
        return True
    if s[0] in _BULLET_CHARS:
        return True
    lower = s.lower()
    for hint in _SECTION_HEADER_HINTS:
        if lower == hint or lower.startswith(hint + " ") or lower.startswith(hint + ":"):
            return True
    # "Company Name | Role | May 2024 – Present" style header
    if re.search(r"\d{4}", s[:80]) and ("|" in s[:80] or "—" in s[:80] or "–" in s[:80]):
        return True
    return False


def _reflow_pdf_text(raw: str) -> str:
    """Reconstruct paragraphs from pypdf's layout-preserving output.

    Two artifacts to fix:
      1. Doubled/tripled spaces inside a line from PDF kerning.
      2. Wrapped-column word-per-line: a paragraph that wrapped inside
         the PDF becomes 8 lines of 1 word each. We JOIN consecutive
         lines when the previous doesn't end with terminal punctuation
         AND the next isn't the start of a new block (bullet, section,
         dated role line).

    Preserves paragraph breaks (blank lines) and block boundaries.
    """
    # 1. Collapse ≥2 spaces per line + strip trailing whitespace.
    # ALSO drop lines that are only whitespace: pypdf inserts single-space
    # "lines" between word-per-line output as layout padding. If we treat
    # those as blank lines we get "paragraph breaks" between every word
    # and never merge. Real paragraph breaks in a resume are extremely
    # rare inside a section, so being greedy about merging is safe.
    lines = [re.sub(r" {2,}", " ", ln.rstrip()) for ln in raw.splitlines()]
    lines = [ln for ln in lines if ln.strip() != ""]  # drop artifacts

    out: list[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            # (unreachable now — we filtered above, kept for clarity)
            if out and out[-1] != "":
                out.append("")
            continue

        # If the previous line exists, has content, and doesn't end
        # with terminal punctuation, AND this line isn't the start
        # of a new block, MERGE with space.
        if (out and out[-1]
                and not out[-1].rstrip().endswith((".", "!", "?", ":", ";", "|"))
                and not _looks_like_new_block(s)):
            out[-1] = out[-1] + " " + s
            continue

        out.append(s)

    # Collapse trailing blank + reduce triple-blanks (defense-in-depth).
    while out and out[-1] == "":
        out.pop()
    return re.sub(r"\n{3,}", "\n\n", "\n".join(out))


def _parse_pdf(data: bytes) -> dict[str, Any]:
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    joined = "\n".join(pages)

    # pypdf's extract_text is layout-preserving: text runs positioned at
    # different X-coords come out as separate "lines", so a wrapped
    # paragraph in a narrow column becomes 1 word per line. It also
    # inserts extra spaces where kerning is non-standard, giving us
    # "Ottawa  ,  Ontario" instead of "Ottawa, Ontario". Reflow fixes
    # both before we hand raw_text to the section splitter.
    reflowed = _reflow_pdf_text(joined)

    # PDFs often render bullets as glyphs ('•', '·', '◦'); strip them so
    # downstream code sees clean lines.
    cleaned_lines = []
    for line in reflowed.splitlines():
        cleaned_lines.append(_strip_leading_bullet(line.strip()))
    raw_text = "\n".join(line for line in cleaned_lines if line)

    parsed = _common_parse(raw_text)
    parsed["source_format"] = "pdf"
    parsed["docx_meta"] = {
        "has_tables": False,
        "has_images": False,
        "table_count": 0,
        "image_count": 0,
    }
    parsed["stats"]["page_estimate"] = len(reader.pages)
    return parsed


def _common_parse(raw_text: str) -> dict[str, Any]:
    """Shared logic for both formats once we have raw text."""
    lines = [ln.rstrip() for ln in raw_text.splitlines()]
    nonempty = [ln for ln in lines if ln.strip()]

    contact = _extract_contact(raw_text, nonempty)
    sections = _split_sections(nonempty)

    word_count = len(re.findall(r"\b\w+\b", raw_text))
    # We strip bullet glyphs in the parsers above, so count differently:
    # a bullet is just a non-empty line in any section other than 'summary'
    # / 'header'. Best-effort, only used for stats display.
    bullet_count = sum(1 for ln in nonempty if 5 < len(ln) < 250)
    page_estimate = max(1, round(word_count / 400))

    return {
        "raw_text": raw_text,
        "contact": contact,
        "sections": sections,
        "stats": {
            "word_count": word_count,
            "page_estimate": page_estimate,
            "bullet_count": bullet_count,
        },
    }


def _extract_contact(raw_text: str, nonempty: list[str]) -> dict[str, str]:
    email_match = EMAIL_RE.search(raw_text)
    phone_match = PHONE_RE.search(raw_text)
    linkedin_match = LINKEDIN_RE.search(raw_text)
    location_match = LOCATION_RE.search(raw_text)

    # Name is the trickiest. Heuristic: first non-empty line that isn't
    # contact info, is 2–5 words, and starts with a capital.
    name = ""
    for line in nonempty[:5]:
        stripped = line.strip()
        if EMAIL_RE.search(stripped) or PHONE_RE.search(stripped):
            continue
        if URL_RE.search(stripped):
            continue
        words = stripped.split()
        # Accept 1-6 words (single-word names exist in many cultures — the
        # 2-word floor rejected valid names like "Mehran" alone). Still
        # reject lines ending in ':' (they're section headers) and lines
        # that look like tag lists or job titles (contain digits).
        if 1 <= len(words) <= 6 and stripped[0].isupper() \
                and not stripped.endswith(":") \
                and not any(c.isdigit() for c in stripped):
            name = stripped
            break

    # Website / portfolio (any URL that isn't linkedin)
    website = ""
    for url in URL_RE.findall(raw_text):
        if "linkedin.com" not in url.lower():
            website = url
            break

    return {
        "name": name,
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "location": location_match.group(0).strip().rstrip(",.") if location_match else "",
        "linkedin": linkedin_match.group(0) if linkedin_match else "",
        "website": website,
    }


def _split_sections(nonempty: list[str]) -> dict[str, list[str]]:
    """Walk the lines and bucket them into sections by header detection."""
    sections: dict[str, list[str]] = {}
    current = "header"   # everything before first detected section goes here
    sections[current] = []

    for line in nonempty:
        matched = _match_section(line)
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)

    # Drop empty sections, except keep 'header' even if empty so callers
    # don't KeyError.
    return {k: v for k, v in sections.items() if v or k == "header"}


def _match_section(line: str) -> str | None:
    stripped = line.strip()
    if len(stripped) > 60:   # section headers are short
        return None
    for name, patterns in SECTION_PATTERNS.items():
        for pat in patterns:
            if re.match(pat, stripped, re.IGNORECASE):
                return name
    return None


_BULLET_GLYPHS = "•◦●·●○◆◇■□▪▫–—-*"


def _strip_leading_bullet(text: str) -> str:
    """Remove a single leading bullet glyph + whitespace from a line."""
    if not text:
        return text
    if text[0] in _BULLET_GLYPHS:
        return text[1:].lstrip(_BULLET_GLYPHS + " \t").strip()
    return text.strip()
