"""Render parsed resume data into a clean, ATS-friendly DOCX.

Design principles:
- Single column (no tables for layout).
- No images, no headers/footers, no text boxes.
- Standard fonts (Calibri 11pt body, 14pt name).
- Section headings are plain bold text, not Word "Heading" styles
  (Word headings sometimes confuse older ATS).
- Bullets use Word's built-in List Bullet style.
- Inside experience, lines that look like job titles (have a year range +
  a separator) are rendered bold instead of as bullets. Achievement
  lines become bullets.
"""
from __future__ import annotations

import io
import re
from typing import Any

from datetime import date

from docx import Document
from docx.shared import Pt, Inches


# A line is treated as a "job title" if it contains a 4-digit year AND
# a separator commonly used in resume titles ('—' '–' '|' '/' ' at ').
_TITLE_LINE_RE = re.compile(
    r"\d{4}.*[—–|/]|[—–|/].*\d{4}|\bat\b.*\d{4}|\(.*\d{4}.*\)",
    re.IGNORECASE,
)


def _looks_like_title(line: str) -> bool:
    return bool(_TITLE_LINE_RE.search(line))


SECTION_ORDER = [
    "summary",
    "experience",
    "education",
    "skills",
    "certifications",
    "projects",
    "publications",
    "volunteer",
    "awards",
    "languages",
    "interests",
    "references",
]

SECTION_TITLES = {
    "summary": "PROFESSIONAL SUMMARY",
    "experience": "EXPERIENCE",
    "education": "EDUCATION",
    "skills": "SKILLS",
    "certifications": "CERTIFICATIONS",
    "projects": "PROJECTS",
    "publications": "PUBLICATIONS",
    "volunteer": "VOLUNTEER EXPERIENCE",
    "awards": "AWARDS",
    "languages": "LANGUAGES",
    "interests": "INTERESTS",
    "references": "REFERENCES",
}


def render_docx(parsed: dict[str, Any]) -> bytes:
    """Build a clean DOCX from a parsed-resume dict. Returns raw bytes."""
    doc = Document()
    _set_default_style(doc)
    _set_narrow_margins(doc)

    contact = dict(parsed.get("contact", {}))
    sections = parsed.get("sections", {})

    # Last-chance name recovery — if the parser (or the LLM tailoring pass)
    # left the name blank, try the first meaningful line of raw_text. Real
    # resumes always have the name at the top; a blank header downloads look
    # broken.
    if not (contact.get("name") or "").strip():
        contact["name"] = _guess_name_from_raw(parsed.get("raw_text") or "")

    _write_header(doc, contact)

    for section_key in SECTION_ORDER:
        items = sections.get(section_key)
        if not items:
            continue
        # 'summary' is prose; everything else is bulleted lists.
        as_bullets = section_key != "summary"
        _write_section(doc, SECTION_TITLES[section_key], items, as_bullets=as_bullets)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def _guess_name_from_raw(raw_text: str) -> str:
    """Fallback name extraction — used when parser + LLM both left it blank.

    Walks the first ~5 non-empty lines and returns the first that looks like
    a person's name: 1-6 words, starts capitalized, no digits, no colons,
    no obvious contact signals (@ signs, common URL substrings).
    """
    if not raw_text:
        return ""
    for line in raw_text.splitlines()[:15]:
        stripped = line.strip()
        if not stripped:
            continue
        if "@" in stripped or "http" in stripped.lower() or "linkedin" in stripped.lower():
            continue
        if any(c.isdigit() for c in stripped):
            continue
        if stripped.endswith(":") or "|" in stripped:
            continue
        words = stripped.split()
        if 1 <= len(words) <= 6 and stripped[0].isupper():
            return stripped
    return ""


def _write_header(doc: Document, contact: dict[str, str]) -> None:
    # If parser couldn't extract a name, skip the name paragraph entirely.
    # Emitting "Your Name" as a placeholder read as unfinished/spammy to the
    # candidate — better to leave blank so they add their own in Word.
    name = (contact.get("name") or "").strip()
    if name:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)

    contact_bits = [
        contact.get("location"),
        contact.get("phone"),
        contact.get("email"),
        contact.get("linkedin"),
        contact.get("website"),
    ]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        para = doc.add_paragraph()
        run = para.add_run(contact_line)
        run.font.size = Pt(10)


def render_cover_letter_docx(
    cover_letter: str,
    contact: dict[str, str],
    company: str = "",
    position: str = "",
) -> bytes:
    """Render a tailored cover letter as a clean DOCX matching the resume's
    look-and-feel (same font, same margins, same header block). Returns raw bytes.

    Layout:
        [Name]                          ← same styling as resume header
        location | phone | email | ...
        (blank)
        August 11, 2026
        (blank)
        {cover letter body — paragraphs preserved from blank-line splits}

    Trusts the LLM's cover_letter to contain its own greeting and closing.
    If the body lacks a signature, we don't add one — better to leave a clean
    end than to fabricate.
    """
    doc = Document()
    _set_default_style(doc)
    _set_narrow_margins(doc)

    # Same name-recovery fallback as render_docx — a blank header on a cover
    # letter reads as unfinished.
    contact = dict(contact or {})
    if not (contact.get("name") or "").strip():
        # We don't have raw_text here; the caller can pre-fill name via contact.
        # But we still fall through cleanly (_write_header skips blank names).
        pass
    _write_header(doc, contact)

    # Date line — spacer paragraph, then the date in body style.
    doc.add_paragraph()
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))

    # Optional "Re:" line if we know what they're applying for.
    if company or position:
        re_bits = [b for b in [position, company] if b]
        re_para = doc.add_paragraph()
        re_run = re_para.add_run("Re: " + " at ".join(re_bits))
        re_run.bold = True

    doc.add_paragraph()

    # Body: split on blank lines to preserve paragraph breaks. Any single
    # newlines inside a paragraph become spaces so wrapping stays natural.
    text = (cover_letter or "").strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    for para_text in paragraphs:
        collapsed = re.sub(r"\s*\n\s*", " ", para_text)
        doc.add_paragraph(collapsed)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _write_section(doc: Document, title: str, items: list[str], *, as_bullets: bool) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    for raw_item in items:
        item = raw_item.strip().lstrip("•◦●·○◆◇■□▪▫–—-* ").strip()
        if not item:
            continue
        # Title-looking lines (e.g. "Sr BIM Coord — Acme (2022 – Present)")
        # render bold, NOT as bullets. Plain achievement lines become bullets.
        if as_bullets and _looks_like_title(item):
            para = doc.add_paragraph()
            run = para.add_run(item)
            run.bold = True
        elif as_bullets:
            doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph(item)
