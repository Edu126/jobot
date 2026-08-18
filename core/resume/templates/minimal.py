"""Minimal — the default template. Matches the current writer.py aesthetic
(Calibri 11pt body, 14–16pt bold headings, all black, no accent color, no
italic) but consumes the STRUCTURED schema — so title / dates split cleanly
across a tab stop instead of relying on regex to bold "title-looking" lines.

Design opinion:
- Single column, always.
- All black. No accent colors.
- No italic anywhere (per user preference).
- Section headings: bold uppercase, 12pt. No borders, no small caps.
- Job entries: bold title left, dates right (tab stop); company · location
  underneath in plain body weight.
- Whitespace: generous but not loud. Matches the feel of the existing
  writer.py output so users can't tell the tailor swapped renderers.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches

from .schema import StructuredResume


BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
NAME_SIZE = Pt(16)
CONTACT_SIZE = Pt(10)
HEADING_SIZE = Pt(12)

PAGE_WIDTH_IN = 8.5
LEFT_MARGIN = 0.7
RIGHT_MARGIN = 0.7
CONTENT_WIDTH = PAGE_WIDTH_IN - LEFT_MARGIN - RIGHT_MARGIN


def render(data: StructuredResume) -> bytes:
    doc = Document()
    _set_base_style(doc)
    _set_margins(doc)

    contact = data.get("contact") or {}
    _write_header(doc, contact)

    if summary := (data.get("summary") or "").strip():
        _write_heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph()
        _set_para_spacing(p, space_after=Pt(6))
        run = p.add_run(summary)
        _style_run(run, size=BODY_SIZE)

    if experience := data.get("experience"):
        _write_heading(doc, "EXPERIENCE")
        for i, entry in enumerate(experience):
            _write_experience_entry(doc, entry, is_last=(i == len(experience) - 1))

    if education := data.get("education"):
        _write_heading(doc, "EDUCATION")
        for entry in education:
            _write_education_entry(doc, entry)

    if skills := data.get("skills"):
        _write_heading(doc, "SKILLS")
        p = doc.add_paragraph()
        _set_para_spacing(p, space_after=Pt(6))
        run = p.add_run(", ".join(str(s).strip() for s in skills if str(s).strip()))
        _style_run(run, size=BODY_SIZE)

    if certs := data.get("certifications"):
        _write_heading(doc, "CERTIFICATIONS")
        for cert in certs:
            _write_cert_entry(doc, cert)

    if projects := data.get("projects"):
        _write_heading(doc, "PROJECTS")
        for proj in projects:
            _write_project_entry(doc, proj)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────
# Header
# ─────────────────────────────────────────────────────────────

def _write_header(doc, contact: dict) -> None:
    name = (contact.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        _set_para_spacing(p, space_before=Pt(0), space_after=Pt(2))
        run = p.add_run(name)
        _style_run(run, size=NAME_SIZE, bold=True)

    bits = [
        contact.get("location"),
        contact.get("phone"),
        contact.get("email"),
        contact.get("linkedin"),
        contact.get("website"),
    ]
    line = " | ".join(b for b in bits if b)
    if line:
        p = doc.add_paragraph()
        _set_para_spacing(p, space_after=Pt(4))
        run = p.add_run(line)
        _style_run(run, size=CONTACT_SIZE)


# ─────────────────────────────────────────────────────────────
# Section heading — small caps + hunter-green + thin bottom rule
# ─────────────────────────────────────────────────────────────

def _write_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=Pt(8), space_after=Pt(2))
    run = p.add_run(text)
    _style_run(run, size=HEADING_SIZE, bold=True)


# ─────────────────────────────────────────────────────────────
# Experience entry — title/dates row, then company row, then bullets
# ─────────────────────────────────────────────────────────────

def _write_experience_entry(doc, entry: dict, is_last: bool) -> None:
    title = (entry.get("title") or "").strip()
    company = (entry.get("company") or "").strip()
    location = (entry.get("location") or "").strip()
    start = (entry.get("start") or "").strip()
    end = (entry.get("end") or "").strip()
    dates = " – ".join(x for x in [start, end] if x)

    # Row 1: bold title left, dates right (via tab stop)
    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=Pt(4), space_after=Pt(0))
    _add_right_tab(p)
    r_title = p.add_run(title)
    _style_run(r_title, size=BODY_SIZE, bold=True)
    if dates:
        p.add_run("\t")
        r_date = p.add_run(dates)
        _style_run(r_date, size=BODY_SIZE)

    # Row 2: company, location (plain body, not italic)
    sub_bits = [b for b in [company, location] if b]
    if sub_bits:
        p2 = doc.add_paragraph()
        _set_para_spacing(p2, space_before=Pt(0), space_after=Pt(2))
        run = p2.add_run(", ".join(sub_bits))
        _style_run(run, size=BODY_SIZE)

    for bullet in entry.get("bullets") or []:
        text = str(bullet).strip().lstrip("•◦●·○◆◇■□▪▫–—-* ").strip()
        if not text:
            continue
        bp = doc.add_paragraph(text, style="List Bullet")
        _set_para_spacing(bp, space_before=Pt(0), space_after=Pt(1))
        for run in bp.runs:
            _style_run(run, size=BODY_SIZE)


# ─────────────────────────────────────────────────────────────
# Education — single line: Degree — School, Location, Year
# ─────────────────────────────────────────────────────────────

def _write_education_entry(doc, entry: dict) -> None:
    degree = (entry.get("degree") or "").strip()
    school = (entry.get("school") or "").strip()
    location = (entry.get("location") or "").strip()
    year = (entry.get("year") or "").strip()

    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=Pt(2), space_after=Pt(2))
    _add_right_tab(p)

    left = degree
    if school:
        left = f"{degree} — {school}" if degree else school
    if location:
        left = f"{left}, {location}" if left else location

    r_left = p.add_run(left)
    _style_run(r_left, size=BODY_SIZE)
    if year:
        p.add_run("\t")
        r_year = p.add_run(year)
        _style_run(r_year, size=BODY_SIZE)

    if notes := (entry.get("notes") or "").strip():
        n = doc.add_paragraph()
        _set_para_spacing(n, space_before=Pt(0), space_after=Pt(2))
        run = n.add_run(notes)
        _style_run(run, size=BODY_SIZE)


# ─────────────────────────────────────────────────────────────
# Cert — single line: Name — Issuer, Year
# ─────────────────────────────────────────────────────────────

def _write_cert_entry(doc, cert: dict) -> None:
    name = (cert.get("name") or "").strip()
    issuer = (cert.get("issuer") or "").strip()
    year = (cert.get("year") or "").strip()

    p = doc.add_paragraph()
    _set_para_spacing(p, space_before=Pt(1), space_after=Pt(1))
    _add_right_tab(p)

    left_parts = [x for x in [name, issuer] if x]
    left = " — ".join(left_parts) if left_parts else ""
    r = p.add_run(left)
    _style_run(r, size=BODY_SIZE)
    if year:
        p.add_run("\t")
        r_y = p.add_run(year)
        _style_run(r_y, size=BODY_SIZE)


# ─────────────────────────────────────────────────────────────
# Project — name row + description + bullets
# ─────────────────────────────────────────────────────────────

def _write_project_entry(doc, proj: dict) -> None:
    name = (proj.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        _set_para_spacing(p, space_before=Pt(2), space_after=Pt(0))
        r = p.add_run(name)
        _style_run(r, size=BODY_SIZE, bold=True)

    if desc := (proj.get("description") or "").strip():
        p = doc.add_paragraph()
        _set_para_spacing(p, space_before=Pt(0), space_after=Pt(2))
        r = p.add_run(desc)
        _style_run(r, size=BODY_SIZE)

    for bullet in proj.get("bullets") or []:
        text = str(bullet).strip()
        if not text:
            continue
        bp = doc.add_paragraph(text, style="List Bullet")
        _set_para_spacing(bp, space_before=Pt(0), space_after=Pt(1))
        for run in bp.runs:
            _style_run(run, size=BODY_SIZE)


# ─────────────────────────────────────────────────────────────
# Low-level docx helpers
# ─────────────────────────────────────────────────────────────

def _set_base_style(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE


def _set_margins(doc) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(LEFT_MARGIN)
        section.right_margin = Inches(RIGHT_MARGIN)


def _set_para_spacing(paragraph, *, space_before=None, space_after=None) -> None:
    pf = paragraph.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after


def _style_run(run, *, size=None, bold=False) -> None:
    """Minimal template keeps things simple: size + optional bold, no color,
    no italic. Everything else picks up the Normal style (Calibri, black)."""
    if size is not None:
        run.font.size = size
    run.font.bold = bold


def _add_right_tab(paragraph) -> None:
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(CONTENT_WIDTH), WD_TAB_ALIGNMENT.RIGHT)
