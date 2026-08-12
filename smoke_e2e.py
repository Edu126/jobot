"""End-to-end audit script.

Simulates a realistic user flow with a real-shape resume and JD,
exercising every module and reporting issues found.
"""
from __future__ import annotations

import io
import json
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document

import core.db as db
from core.llm.rewrite import rewrite_resume
from core.matching.tfidf_match import match as tfidf_match
from core.resume.ats import run_checks
from core.resume.parser import parse_resume
from core.resume.section_presence import analyze as section_analyze
from core.resume.writer import render_docx


# ---------- realistic fixtures ----------

def make_realistic_resume_docx() -> bytes:
    """Resume styled like real ones — non-standard headings, mixed bullet
    glyphs, hard-coded contact line, 350 words."""
    d = Document()
    d.add_paragraph("Carlos Mendez")
    d.add_paragraph(
        "Ottawa, Ontario, Canada  |  613-555-9999  |  "
        "carlos.mendez@example.com  |  linkedin.com/in/carlosmendez"
    )

    d.add_paragraph("Career Profile")
    d.add_paragraph(
        "BIM Coordinator and Estimator with 6 years in mid-rise residential "
        "and commercial construction across Ottawa and Gatineau. Experienced "
        "in Revit, Navisworks, AutoCAD, and Bluebeam. Bilingual (English fluent, "
        "French intermediate)."
    )

    d.add_paragraph("Professional Background")
    d.add_paragraph("Senior BIM Coordinator — Tessier Construction (March 2022 – Present)")
    d.add_paragraph("Led BIM coordination on a $42M commercial office tower in downtown Ottawa", style="List Bullet")
    d.add_paragraph("Reduced RFIs by 35% by implementing weekly clash detection with Navisworks", style="List Bullet")
    d.add_paragraph("Coordinated Revit models across structural, MEP, and architectural disciplines", style="List Bullet")
    d.add_paragraph("Trained 4 junior modelers on BIM 360 workflows and Bluebeam markup standards", style="List Bullet")

    d.add_paragraph("BIM Modeler / Junior Estimator — Beauchamp Builders (Jan 2020 – Feb 2022)")
    d.add_paragraph("Developed Revit families for 14 residential projects (40-200 units each)", style="List Bullet")
    d.add_paragraph("Produced quantity takeoffs in Bluebeam, contributing to bids worth $18M total", style="List Bullet")
    d.add_paragraph("Coordinated weekly with project managers on schedule changes using MS Project", style="List Bullet")

    d.add_paragraph("BIM Drafter — Lavoie Architects (Sept 2018 – Dec 2019)")
    d.add_paragraph("Drafted construction documents in AutoCAD and Revit for residential renovations", style="List Bullet")
    d.add_paragraph("Performed initial cost estimates for design-build packages", style="List Bullet")

    d.add_paragraph("Academic Credentials")
    d.add_paragraph(
        "Architectural Technology Diploma — Algonquin College, Ottawa (2018)"
    )
    d.add_paragraph(
        "Bachelor of Civil Engineering — Universidad Nacional de Colombia (2015)"
    )

    d.add_paragraph("Software & Tools")
    d.add_paragraph(
        "Revit · Navisworks · AutoCAD · Civil 3D · BIM 360 · Bluebeam · "
        "MS Project · Procore · Microsoft Office · Python (basic)"
    )

    d.add_paragraph("Certifications")
    d.add_paragraph("LEED Green Associate — USGBC, 2023")
    d.add_paragraph("WHMIS 2015 — Government of Canada, 2022")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


REAL_JD = """
BIM Coordinator — Construction (Ottawa, ON)

Tessier Construction is seeking an experienced BIM Coordinator to lead model coordination on a
new mixed-use development in downtown Ottawa. The successful candidate will be responsible for
delivering coordinated Revit models across architectural, structural, and MEP disciplines, and
running weekly clash detection meetings using Navisworks.

Requirements:
- 5+ years experience as a BIM Coordinator on commercial or mixed-use projects
- Advanced Revit, Navisworks, BIM 360, and AutoCAD skills
- Knowledge of IFC, COBie, and LOD 350 standards
- Familiarity with quantity takeoffs (Bluebeam preferred)
- Experience training junior modelers
- Bilingual (English/French) is an asset but not required
- LEED accreditation preferred
- Diploma or degree in Architectural Technology, Civil Engineering, or related field

Responsibilities:
- Lead BIM coordination meetings with subcontractors and design teams
- Maintain and audit federated Revit models
- Establish and enforce BIM execution plan standards
- Mentor junior BIM staff
- Coordinate with project managers on schedule integration via MS Project
"""


# ---------- runner ----------

bugs: list[str] = []
warnings: list[str] = []
findings: list[str] = []


def bug(msg: str):
    print(f"  🔴 BUG: {msg}")
    bugs.append(msg)


def warn(msg: str):
    print(f"  🟡 WARN: {msg}")
    warnings.append(msg)


def ok(msg: str):
    print(f"  🟢 {msg}")
    findings.append(msg)


def section(title: str):
    print()
    print("=" * 70)
    print(f" {title}")
    print("=" * 70)


# ---------- audit steps ----------

def audit_parse(resume_bytes: bytes) -> dict:
    section("STEP 1: PARSE")
    parsed = parse_resume(resume_bytes, "carlos.docx")

    found_sections = list(parsed["sections"].keys())
    print(f"  sections found: {found_sections}")
    print(f"  contact: {parsed['contact']}")
    print(f"  stats: {parsed['stats']}")

    # Contact checks
    if not parsed["contact"]["email"]:
        bug("email not extracted")
    else:
        ok(f"email extracted: {parsed['contact']['email']}")
    if not parsed["contact"]["phone"]:
        bug("phone not extracted")
    else:
        ok(f"phone extracted: {parsed['contact']['phone']}")
    if not parsed["contact"]["linkedin"]:
        bug("linkedin not extracted")
    else:
        ok(f"linkedin extracted: {parsed['contact']['linkedin']}")
    if not parsed["contact"]["name"]:
        bug("name not extracted")
    else:
        ok(f"name extracted: {parsed['contact']['name']}")

    # Section bucketing — expanded synonyms should now catch these
    expected = ["summary", "experience", "education", "skills", "certifications"]
    missing = [s for s in expected if s not in found_sections]
    if missing:
        bug(f"sections NOT bucketed via heading: {missing}")
    else:
        ok("all expected sections bucketed via heading")

    # Bullet contents — should be clean (no leading '•')
    for sect_name, items in parsed["sections"].items():
        for item in items:
            if item.startswith(("•", "·", "◦", "●")):
                bug(f"section '{sect_name}' has bullet glyph in text: {item[:40]!r}")
                break

    return parsed


def audit_ats(parsed: dict):
    section("STEP 2: ATS REPORT")
    report = run_checks(parsed)
    print(f"  score: {report['score']}/100")
    print(f"  issues: {len(report['issues'])}, passed: {len(report['passed'])}")
    for i in report["issues"]:
        print(f"    [{i['severity']}] {i['message'][:90]}")

    if report["score"] < 80:
        bug(f"realistic resume scored only {report['score']}/100 — false positives likely")
    else:
        ok(f"score is {report['score']}/100 — reasonable for this resume")

    # Should NOT flag experience/education/skills as missing
    bad_msgs = [
        i for i in report["issues"]
        if i["severity"] in ("critical", "warning")
        and ("No Experience" in i["message"] or "No Education" in i["message"] or "No Skills" in i["message"])
    ]
    if bad_msgs:
        bug(f"false-positive missing-section flags: {[i['message'][:60] for i in bad_msgs]}")
    else:
        ok("no false-positive missing-section flags")


def audit_tfidf(parsed: dict):
    section("STEP 3: TF-IDF MATCH")
    result = tfidf_match(parsed["raw_text"], REAL_JD)
    print(f"  similarity score: {result['similarity_score']}/100")
    print(f"  matched (top 15): {result['matched_keywords'][:15]}")
    print(f"  missing: {result['missing_keywords']}")

    if result["similarity_score"] < 40:
        bug(f"score only {result['similarity_score']}/100 — should be high; this resume directly matches the JD")
    else:
        ok(f"match score {result['similarity_score']}/100 — sensible")

    # Bug check: terms that are clearly IN the resume should not be in missing
    resume_lower = parsed["raw_text"].lower()
    false_missing = []
    for term in result["missing_keywords"]:
        # unigram check: every word of the term appears in the resume
        words = term.split()
        if all(w in resume_lower for w in words):
            false_missing.append(term)
    if false_missing:
        bug(f"missing-keywords contains terms whose words all appear in resume: {false_missing}")
    else:
        ok("missing-keywords are genuinely absent from resume")

    # Bug check: generic noise words that should be filtered
    noise_in_missing = [t for t in result["missing_keywords"] if t in {"experience", "education", "background", "years", "ability"}]
    if noise_in_missing:
        bug(f"generic noise words in missing list: {noise_in_missing}")


def audit_section_presence(parsed: dict):
    section("STEP 4: SECTION PRESENCE (content-based)")
    ev = section_analyze(parsed)
    for sect, e in ev.items():
        emoji = "✅" if e.present else "❌"
        print(f"  {emoji} {sect:15s}: via={e.via}, note={e.note}")
    for sect in ("experience", "education", "skills", "certifications"):
        if not ev[sect].present:
            bug(f"section_presence failed to detect '{sect}' in realistic resume")


def audit_tailor(parsed: dict) -> dict:
    section("STEP 5: TAILORING (stub LLM)")

    # Capture what we send to the LLM
    captured_prompt = {}

    class StubClient:
        def generate_json(self, prompt, **kw):
            captured_prompt["prompt"] = prompt
            return {
                "sections": {
                    "summary": ["Tailored summary mentioning IFC, LOD 350, and clash detection."],
                    "experience": [
                        "Senior BIM Coordinator — Tessier Construction (March 2022 – Present)",
                        "Led BIM coordination for $42M mixed-use development, enforcing LOD 350 standards.",
                        "Established weekly clash detection cadence in Navisworks, reducing RFIs 35%.",
                    ],
                    "education": [
                        "Architectural Technology Diploma — Algonquin College, Ottawa (2018)",
                    ],
                    "skills": ["Revit · Navisworks · AutoCAD · BIM 360 · IFC · COBie · Bluebeam"],
                    "certifications": ["LEED Green Associate — USGBC, 2023"],
                },
                "cover_letter": "Dear Hiring Manager,\n\nI am writing to apply for the BIM Coordinator role at Tessier Construction...\n\nSincerely,\nCarlos Mendez",
                "notes": "Surfaced IFC, COBie, LOD 350 — all anchored to actual experience.",
                "warnings": [],
            }

    tailored = rewrite_resume(parsed, REAL_JD, "balanced", StubClient())

    # Verify what was sent — should include all the parsed sections, not just empty
    sent = captured_prompt["prompt"]
    if '"experience":' not in sent:
        bug("LLM prompt did not include 'experience' section — was the parser empty?")
    elif '"Senior BIM Coordinator"' in sent or "Senior BIM Coordinator" in sent:
        ok("LLM received actual experience content")
    else:
        warn("LLM prompt has 'experience' key but content unclear")

    # Output schema checks
    if not tailored.get("cover_letter"):
        bug("cover letter missing from tailored output")
    elif "Dear Hiring Manager" in tailored["cover_letter"]:
        ok("cover letter present")
    if tailored.get("tailoring_level") != "balanced":
        bug(f"tailoring_level mismatch: {tailored.get('tailoring_level')!r}")
    if "experience" not in tailored.get("sections", {}):
        bug("tailored output missing experience section")

    # Render to DOCX
    docx_bytes = render_docx(tailored)
    if docx_bytes[:2] != b"PK":
        bug("tailored DOCX is not a valid zip")
    else:
        ok(f"tailored DOCX renders ({len(docx_bytes)} bytes)")

    # Check bullets render as bullets in the output DOCX
    out_doc = Document(io.BytesIO(docx_bytes))
    bullet_count = sum(
        1 for p in out_doc.paragraphs
        if "list bullet" in (p.style.name or "").lower()
    )
    if bullet_count < 2:
        bug(f"DOCX has only {bullet_count} bullet paragraphs — expected ~4")
    else:
        ok(f"DOCX has {bullet_count} bullet paragraphs")

    # Title-detection check: 'Senior BIM Coordinator — Tessier ... (2022 – Present)'
    # should be a bold paragraph, NOT a bullet.
    title_text = "Senior BIM Coordinator"
    title_paras = [p for p in out_doc.paragraphs if title_text in p.text]
    if not title_paras:
        bug("title line missing from output DOCX entirely")
    else:
        tp = title_paras[0]
        if "list bullet" in (tp.style.name or "").lower():
            bug("job-title line rendered as a BULLET — should be bold paragraph")
        else:
            ok("job-title line rendered as non-bullet paragraph")
        if not any(r.bold for r in tp.runs):
            warn("job-title line not bold (acceptable but less ideal)")
        else:
            ok("job-title line is bold")

    return tailored


def audit_db(parsed: dict, tailored: dict, raw_bytes: bytes):
    section("STEP 6: DATABASE ROUND-TRIP")
    tmp = Path(tempfile.mkstemp(suffix=".db")[1])
    try:
        db.init_db(tmp)
        rid = db.save_resume("carlos.docx", parsed, raw_bytes, set_current=True, path=tmp)
        ok(f"resume saved: id={rid}")

        loaded = db.get_current_resume(path=tmp)
        if loaded["parsed"]["contact"] != parsed["contact"]:
            bug("contact mismatch after DB round-trip")
        else:
            ok("contact survived round-trip")
        if loaded["parsed"]["sections"] != parsed["sections"]:
            bug("sections mismatch after DB round-trip")
        else:
            ok("sections survived round-trip")

        # Save a fake scraped job
        job = {
            "id": "test-job-1", "title": "BIM Coordinator", "company": "Tessier Construction",
            "location": "Ottawa, ON", "site": "linkedin",
            "date_posted": "2026-06-15", "job_url": "https://example/j1",
            "description": REAL_JD, "is_remote": False,
            "min_salary": 75000, "max_salary": 105000,
            "detected_language": "en", "french_required": False,
        }
        db.upsert_job(job, path=tmp)
        app_id = db.create_or_get_application("test-job-1", status="interested", resume_id=rid, path=tmp)
        db.update_application(
            app_id,
            status="applied",
            tailoring_level=tailored.get("tailoring_level"),
            tailored_resume_json=json.dumps({"sections": tailored["sections"]}, ensure_ascii=False),
            tailored_cover_letter=tailored.get("cover_letter"),
            path=tmp,
        )
        a = db.get_application(app_id, path=tmp)
        if a["status"] != "applied":
            bug("status not persisted")
        if not a["applied_at"]:
            bug("applied_at not set")
        if "Dear Hiring Manager" not in (a["tailored_cover_letter"] or ""):
            bug("cover letter not persisted")
        else:
            ok("application + tailored content persisted")

        listed = db.list_applications(statuses=["applied"], path=tmp)
        if len(listed) != 1:
            bug(f"list_applications returned {len(listed)} apps, expected 1")
        elif listed[0].get("job_title") != "BIM Coordinator":
            bug("joined job_title missing in list_applications")
        else:
            ok("list_applications joined view works")

    finally:
        tmp.unlink(missing_ok=True)


def audit_ui_compiles():
    section("STEP 7: WEB APP COMPILES")
    import py_compile
    for target in (
        "ui_web/main.py",
        "ui_web/routes/jobs.py",
        "ui_web/routes/applications.py",
        "ui_web/routes/profile.py",
    ):
        try:
            py_compile.compile(target, doraise=True)
            ok(f"{target} compiles cleanly")
        except py_compile.PyCompileError as e:
            bug(f"{target} compile error: {e}")


# ---------- main ----------

def main():
    print("End-to-end audit of the jobot-app build.\n")

    raw_bytes = make_realistic_resume_docx()
    parsed = audit_parse(raw_bytes)
    audit_section_presence(parsed)
    audit_ats(parsed)
    audit_tfidf(parsed)
    tailored = audit_tailor(parsed)
    audit_db(parsed, tailored, raw_bytes)
    audit_ui_compiles()

    section("VERDICT")
    print(f"Findings: {len(findings)}")
    print(f"Warnings: {len(warnings)}")
    print(f"BUGS:     {len(bugs)}")
    if bugs:
        print("\nBug list:")
        for b in bugs:
            print(f"  🔴 {b}")
    print()


if __name__ == "__main__":
    main()
